#!/usr/bin/env python3
"""
document_renderer.py — CPSWC v0 Formal Table Renderer (Step 13B-1)

从 RuntimeSnapshot / FrozenSubmissionInput / calculator_results 取数,
生成正式 DOCX 文档。

本轮只渲染 3 个输出件:
  1. 项目运行摘要页
  2. AT-02 加权综合防治指标计算表
  3. 补偿费计费表

设计边界:
  - 只从 snapshot / frozen / calculator_results 取数, 不回读 sample
  - 不做整本报告书 / 封面 / 目录 / 正文
  - 不做 PDF
  - 不做 Word 模板大系统
  - 表格格式: 中文行业惯例 (表头灰底加粗, 文字左对齐, 数字右对齐, 带标题和来源注脚)

使用方式:
  from cpswc.renderers.document import render_formal_tables
  docx_paths = render_formal_tables(snapshot_dict, frozen_dict, calc_results_dir, output_dir)
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import yaml
from cpswc.paths import GOVERNANCE_DIR

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ============================================================
# Styling helpers
# ============================================================

_HEADER_BG = "D9D9D9"  # 灰色表头背景
_FONT_NAME = "SimSun"   # 宋体 (中文行业标准字体)
_FONT_NAME_FALLBACK = "Arial"  # 英文回退


def _set_cell_shading(cell, color_hex: str):
    """设置单元格背景色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elem)


def _style_header_row(row):
    """表头行: 灰底加粗居中"""
    for cell in row.cells:
        _set_cell_shading(cell, _HEADER_BG)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)


def _add_cell_text(cell, text: str, align: str = "left", bold: bool = False):
    """向单元格写入文本并设置对齐"""
    p = cell.paragraphs[0]
    if align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.size = Pt(9)
    run.bold = bold


def _add_heading(doc: Document, text: str, level: int = 1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.size = Pt(14 if level == 1 else 12)


def _add_footnote(doc: Document, text: str):
    """添加来源注脚"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True


def _add_kv_table(doc: Document, pairs: list[tuple[str, str]]):
    """添加 key-value 两列表格"""
    table = doc.add_table(rows=len(pairs), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (k, v) in enumerate(pairs):
        _add_cell_text(table.rows[i].cells[0], k, bold=True)
        _add_cell_text(table.rows[i].cells[1], str(v))
    return table


# ============================================================
# 1. 项目运行摘要页
# ============================================================

def _render_summary(doc: Document, snapshot: dict, frozen: dict | None):
    """渲染项目运行摘要页"""
    _add_heading(doc, "项目运行摘要", level=1)

    summary = snapshot.get("project_input_summary") or {}
    manifest = snapshot.get("manifest") or {}

    _add_heading(doc, "基本信息", level=2)
    _add_kv_table(doc, [
        ("项目名称", summary.get("name", "")),
        ("项目代码", summary.get("code", "")),
        ("行业分类", str(summary.get("industry", ""))),
        ("方案类型", summary.get("species", "")),
        ("规则集", snapshot.get("ruleset", "")),
        ("生命周期", snapshot.get("lifecycle", "")),
    ])

    _add_heading(doc, "运行结果概览", level=2)
    _add_kv_table(doc, [
        ("Snapshot ID", snapshot.get("snapshot_id", "")),
        ("事实字段数", str(snapshot.get("facts_count", 0))),
        ("派生字段数", str(len(snapshot.get("derived_fields") or {}))),
        ("Live Calculators", str(len(snapshot.get("calculator_results") or []))),
        ("触发义务数", str(len(snapshot.get("triggered_obligations") or []))),
        ("未触发义务数", str(len(snapshot.get("not_triggered_obligations") or []))),
        ("所需制品数", str(len(snapshot.get("required_artifacts") or []))),
        ("所需保障数", str(len(snapshot.get("required_assurances") or []))),
    ])

    if frozen:
        _add_heading(doc, "冻结快照", level=2)
        _add_kv_table(doc, [
            ("Content Hash (SHA256)", frozen.get("content_hash", "")),
            ("Fact Snapshot Hash", frozen.get("fact_snapshot_hash", "")),
            ("冻结时间", frozen.get("frozen_at", "")),
            ("Artifact Manifest", f"{len(frozen.get('artifact_manifest') or [])} items"),
            ("Calculator Manifest", ", ".join(frozen.get("calculator_manifest") or [])),
        ])

    # Calculator 摘要表
    calc_results = snapshot.get("calculator_results") or []
    if calc_results:
        _add_heading(doc, "Calculator 执行摘要", level=2)
        table = doc.add_table(rows=1 + len(calc_results), cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        headers = ["Calculator ID", "输出字段", "状态", "结果"]
        for i, h in enumerate(headers):
            _add_cell_text(table.rows[0].cells[i], h, align="center", bold=True)
        _style_header_row(table.rows[0])
        for j, cr in enumerate(calc_results):
            row = table.rows[j + 1]
            _add_cell_text(row.cells[0], cr.get("calculator_id", ""))
            _add_cell_text(row.cells[1], cr.get("output_field_id", ""))
            _add_cell_text(row.cells[2], cr.get("status", ""), align="center")
            val = cr.get("value")
            if isinstance(val, (dict, list)):
                val_str = f"({type(val).__name__})"
            else:
                val_str = f"{val} {cr.get('unit', '')}"
            _add_cell_text(row.cells[3], val_str, align="right")

    _add_footnote(doc, f"Generated by CPSWC v0 Runtime Service | {snapshot.get('timestamp', '')}")


# ============================================================
# 2. AT-02 加权综合防治指标计算表
# ============================================================

def _render_weighted_target_table(doc: Document, snapshot: dict):
    """渲染 AT-02 加权综合防治指标计算表"""
    doc.add_page_break()
    _add_heading(doc, "附表 AT-02 水土流失防治指标计算表", level=1)

    derived = snapshot.get("derived_fields") or {}
    wt = derived.get("field.derived.target.weighted_comprehensive_target")

    if wt is None:
        doc.add_paragraph("(本项目未触发加权综合目标计算)")
        return

    if isinstance(wt, dict):
        # Record 型: 6 项指标
        _add_heading(doc, "加权综合防治指标值", level=2)

        INDICATOR_LABELS = {
            "control_degree": ("水土流失治理度", "%"),
            "soil_loss_control_ratio": ("土壤流失控制比", ""),
            "spoil_protection_rate": ("渣土防护率", "%"),
            "topsoil_protection_rate": ("表土保护率", "%"),
            "vegetation_restoration_rate": ("林草植被恢复率", "%"),
            "vegetation_coverage_rate": ("林草覆盖率", "%"),
        }

        table = doc.add_table(rows=1 + len(INDICATOR_LABELS), cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for i, h in enumerate(["防治指标", "加权目标值", "单位"]):
            _add_cell_text(table.rows[0].cells[i], h, align="center", bold=True)
        _style_header_row(table.rows[0])

        for j, (key, (label, unit)) in enumerate(INDICATOR_LABELS.items()):
            row = table.rows[j + 1]
            _add_cell_text(row.cells[0], label)
            val = wt.get(key, "")
            _add_cell_text(row.cells[1], str(val), align="right")
            _add_cell_text(row.cells[2], unit, align="center")

        _add_footnote(doc,
            "数据来源: cal.target.weighted_comprehensive | "
            "依据: GB/T 50434-2018 表 4.0.2-5 南方红壤区 | "
            "方法: 按防治分区面积加权")

    else:
        doc.add_paragraph(f"(加权目标为非 record 类型: {type(wt).__name__})")


# ============================================================
# 3. 补偿费计费表
# ============================================================

def _render_compensation_fee_table(doc: Document, snapshot: dict, calc_results_dir: Path | None):
    """渲染水土保持补偿费计费表"""
    doc.add_page_break()
    _add_heading(doc, "水土保持补偿费计费表", level=1)

    derived = snapshot.get("derived_fields") or {}
    fee_amount = derived.get("field.derived.investment.compensation_fee_amount")

    # 尝试从 calculator_results 目录读取详细中间量
    fee_detail = None
    if calc_results_dir and (calc_results_dir / "cal_compensation_fee.json").exists():
        with (calc_results_dir / "cal_compensation_fee.json").open() as f:
            fee_detail = json.load(f)

    # 主表
    _add_heading(doc, "计算参数与结果", level=2)

    # 从 snapshot 的 calculator_results 里找补偿费 calculator
    calc_results = snapshot.get("calculator_results") or []
    comp_calc = next((cr for cr in calc_results
                      if cr.get("calculator_id") == "cal.compensation.fee"
                      and cr.get("status") == "ok"), None)

    if comp_calc is None and fee_amount is None:
        doc.add_paragraph("(本项目无补偿费计算结果)")
        return

    # 单张两列表: 参数/公式/输出/结果/状态/来源
    # (不混两张表, 不保留空白占位行)
    info_rows = [
        ("计算器", "cal.compensation.fee"),
        ("计算公式", "(永久占地 + 临时占地) x 10000 x 费率 / 10000"),
        ("输出字段", "field.derived.investment.compensation_fee_amount"),
        ("应缴补偿费", f"{fee_amount} 万元" if fee_amount is not None else "N/A"),
        ("状态", comp_calc.get("status", "N/A") if comp_calc else "N/A"),
        ("费率依据", "粤发改价格〔2021〕231 号 (0.6 元/m2, 广东口径含临时)"),
    ]

    table = doc.add_table(rows=1 + len(info_rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _add_cell_text(table.rows[0].cells[0], "项目", align="center", bold=True)
    _add_cell_text(table.rows[0].cells[1], "内容", align="center", bold=True)
    _style_header_row(table.rows[0])
    for i, (k, v) in enumerate(info_rows):
        _add_cell_text(table.rows[i + 1].cells[0], k)
        _add_cell_text(table.rows[i + 1].cells[1], str(v), align="right")

    _add_footnote(doc,
        "数据来源: cal.compensation.fee | "
        "费率依据: 粤发改价格〔2021〕231 号 | "
        "计征范围: 永久占地 + 临时占地 (广东口径)")


# ============================================================
# Step 13B-2: NarrativeBook Skeleton
# ============================================================
# ============================================================
# DisplayNumberingPolicy_v0: 从 YAML config 加载章节树
# 宪法必做项 #6: renderer 不再硬编码章节号树
# ============================================================

def _load_chapter_tree() -> list:
    """从 DisplayNumberingPolicy_v0.yaml 加载章节树。"""
    policy_path = GOVERNANCE_DIR / "DisplayNumberingPolicy_v0.yaml"
    if not policy_path.exists():
        raise FileNotFoundError(
            f"DisplayNumberingPolicy_v0.yaml not found: {policy_path}")
    with policy_path.open(encoding="utf-8") as f:
        data = yaml.safe_load(f)
    return data.get("chapter_tree") or []


# 延迟加载, 模块级缓存
_CHAPTER_TREE: list | None = None


def _get_chapter_tree() -> list:
    global _CHAPTER_TREE
    if _CHAPTER_TREE is None:
        _CHAPTER_TREE = _load_chapter_tree()
    return _CHAPTER_TREE


def _try_render_section_table(doc: Document, section_id: str, snapshot: dict):
    """检查某 section 是否有对应的自动生成表, 如有则渲染"""
    try:
        from cpswc.renderers.table_projections import TABLE_PROJECTIONS
        from cpswc.renderers.table_protocol import render_data_table
    except ImportError:
        return
    # 查找 section_id 对应的表 (通过 spec.section_id 匹配)
    for table_id, proj_fn in TABLE_PROJECTIONS.items():
        spec = proj_fn.__code__.co_consts  # 不够可靠, 改用 spec 属性
    # 更可靠的方式: 遍历所有 projection 的 spec, 看哪个 section_id 匹配
    from cpswc.renderers import table_projections as tp
    for attr_name in dir(tp):
        if attr_name.startswith("SPEC_"):
            spec = getattr(tp, attr_name)
            if hasattr(spec, "section_id") and spec.section_id == section_id:
                proj_fn = TABLE_PROJECTIONS.get(spec.table_id)
                if proj_fn:
                    table_data = proj_fn(snapshot)
                    if table_data.rows:  # 有数据才渲染
                        render_data_table(doc, table_data)


def _is_section_active(node: dict, triggered: set[str]) -> bool:
    """判断节点是否活跃 (义务触发或 always)"""
    policy = node.get("render_policy", "always")
    if policy == "always":
        return True
    ob = node.get("conditional_on_obligation")
    if ob and ob in triggered:
        return True
    return False


def _render_section_node(doc: Document, node: dict, triggered: set[str],
                         snapshot: dict, calc_results_dir: Path | None,
                         depth: int = 1,
                         narrative_lookup: dict | None = None):
    """递归渲染章节树节点, 优先消费 NarrativeBlock"""
    policy = node.get("render_policy", "always")
    active = _is_section_active(node, triggered)
    display_num = node.get("display_number", "")
    title = node.get("display_title", "")
    heading_text = f"{display_num} {title}"
    sec_id = node.get("stable_id", "")

    if policy == "omit_if_false" and not active:
        return

    heading_level = min(depth, 3)
    _add_heading(doc, heading_text, level=heading_level)

    # Check for NarrativeBlock (Step 14-0 narrative projection)
    nb = (narrative_lookup or {}).get(sec_id)
    if nb is not None:
        rs = nb.render_status if hasattr(nb, "render_status") else nb.get("render_status")
        rs_val = rs.value if hasattr(rs, "value") else str(rs)

        if rs_val == "not_applicable":
            p = doc.add_paragraph()
            run = p.add_run("本项目不涉及此项内容。")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.italic = True
            # Still recurse children (they may have their own blocks)
            for child in node.get("children") or []:
                _render_section_node(doc, child, triggered, snapshot,
                                     calc_results_dir, depth + 1, narrative_lookup)
            return

        if rs_val == "full":
            paragraphs = nb.paragraphs if hasattr(nb, "paragraphs") else nb.get("paragraphs", [])
            for para in paragraphs:
                text = para.text if hasattr(para, "text") else para.get("text", "")
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(10)
                # Add subtle evidence trace as footnote-style
                ev = para.evidence_refs if hasattr(para, "evidence_refs") else para.get("evidence_refs", [])
                if ev:
                    trace_p = doc.add_paragraph()
                    trace_run = trace_p.add_run(
                        f"[trace: {', '.join(ev[:5])}{'...' if len(ev) > 5 else ''}]")
                    trace_run.font.size = Pt(7)
                    trace_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

            # Auto-embed: check if this section has a registered table projection
            _try_render_section_table(doc, sec_id, snapshot)

            # Recurse children
            for child in node.get("children") or []:
                _render_section_node(doc, child, triggered, snapshot,
                                     calc_results_dir, depth + 1, narrative_lookup)
            return

    # Fallback: original skeleton behavior
    if not active and policy == "shell_if_false":
        p = doc.add_paragraph()
        run = p.add_run("本项目不涉及此项内容。")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True
        return

    placeholder = node.get("placeholder_text")
    if placeholder:
        p = doc.add_paragraph()
        run = p.add_run(placeholder)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Auto-embed for skeleton sections too
    _try_render_section_table(doc, sec_id, snapshot)

    for child in node.get("children") or []:
        _render_section_node(doc, child, triggered, snapshot,
                             calc_results_dir, depth + 1, narrative_lookup)


def render_narrative_skeleton(
    snapshot: dict,
    frozen: dict | None = None,
    calc_results_dir: Path | None = None,
    output_dir: Path | str = ".",
    narrative_blocks: list | None = None,
) -> Path:
    """
    生成骨架版报告书 DOCX (Step 13B-2)。

    固定编号 + 条件显隐/留壳:
      - always: 永远渲染
      - shell_if_false: 未触发时保留编号+标题, 正文写"本项目不涉及"
      - omit_if_false: 未触发时整节不输出, 后续编号不补位

    返回: DOCX 文件路径
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    triggered = set(snapshot.get("triggered_obligations") or [])

    # Step 14-0: 运行 narrative projection (如果可用)
    narrative_lookup: dict = {}
    if narrative_blocks:
        for nb in narrative_blocks:
            sid = nb.section_id if hasattr(nb, "section_id") else nb.get("section_id")
            if sid:
                narrative_lookup[sid] = nb
    else:
        try:
            from cpswc.narrative.projection import project_narrative  # type: ignore
            result = project_narrative(snapshot)
            for nb in result.blocks:
                narrative_lookup[nb.section_id] = nb
        except Exception:
            pass  # fallback to skeleton-only

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)

    # 封面 (极简, 不做精修)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("\n\n\n")
    summary = snapshot.get("project_input_summary") or {}
    project_name = summary.get("name", "（项目名称）")
    run = title_p.add_run(project_name)
    run.font.size = Pt(18)
    run.bold = True
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("水土保持方案报告书")
    run.font.size = Pt(16)
    run.bold = True
    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_p.add_run(f"（骨架版 v0 · {snapshot.get('ruleset', '')} · {snapshot.get('timestamp', '')[:10]}）")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # 遍历章节树 (narrative_lookup 优先于 placeholder)
    for chapter in _get_chapter_tree():
        _render_section_node(doc, chapter, triggered, snapshot,
                             calc_results_dir, depth=1,
                             narrative_lookup=narrative_lookup)

    # 附表说明
    doc.add_page_break()
    _add_heading(doc, "附表", level=1)
    _add_footnote(doc,
        "本骨架版已在相关章节内自动嵌入当前可生成的正文表；"
        "其余附表和专项估算表详见 formal_tables_v0.docx。")
    _add_footnote(doc,
        f"Snapshot: {snapshot.get('snapshot_id', '')} | "
        f"Version: {(frozen or {}).get('content_hash', 'N/A')[:16]}...")

    docx_path = output_dir / "narrative_skeleton_v0.docx"
    doc.save(str(docx_path))
    return docx_path


# ============================================================
# Public API
# ============================================================

def render_formal_tables(
    snapshot: dict,
    frozen: dict | None = None,
    calc_results_dir: Path | None = None,
    output_dir: Path | str = ".",
) -> list[Path]:
    """
    生成正式 DOCX 文档。

    参数:
        snapshot: RuntimeSnapshot dict
        frozen: FrozenSubmissionInput dict (可选, 用于摘要页)
        calc_results_dir: package 的 calculator_results/ 目录 (可选)
        output_dir: 输出目录

    返回:
        生成的 DOCX 文件路径列表
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 设置默认字体 (中文环境下 SimSun/宋体)
    style = doc.styles["Normal"]
    font = style.font
    font.size = Pt(10)

    # 渲染 3 个内容块
    _render_summary(doc, snapshot, frozen)
    _render_weighted_target_table(doc, snapshot)
    _render_compensation_fee_table(doc, snapshot, calc_results_dir)

    # 保存
    docx_path = output_dir / "formal_tables_v0.docx"
    doc.save(str(docx_path))

    return [docx_path]


# ============================================================
# CLI
# ============================================================

def _cli() -> int:
    import argparse
    parser = argparse.ArgumentParser(
        description="CPSWC v0 Formal Table Renderer (Step 13B-1)")
    parser.add_argument("snapshot_json",
                        help="RuntimeSnapshot JSON 文件路径")
    parser.add_argument("--frozen", default=None,
                        help="FrozenSubmissionInput JSON 文件路径")
    parser.add_argument("--calc-results-dir", default=None,
                        help="calculator_results/ 目录路径")
    parser.add_argument("-o", "--output", default=".",
                        help="输出目录 (默认当前目录)")
    args = parser.parse_args()

    with open(args.snapshot_json, encoding="utf-8") as f:
        snapshot = json.load(f)

    frozen = None
    if args.frozen:
        with open(args.frozen, encoding="utf-8") as f:
            frozen = json.load(f)

    calc_dir = Path(args.calc_results_dir) if args.calc_results_dir else None

    paths = render_formal_tables(snapshot, frozen, calc_dir, args.output)
    for p in paths:
        print(f"DOCX written: {p}")
    return 0


if __name__ == "__main__":
    import sys
    sys.exit(_cli())
