"""
table_protocol.py — CPSWC P0 表格生成协议

定义可复用的表格数据结构和通用 DOCX 渲染函数。
所有 CAN_GENERATE 表格共用此协议, 每张表只需提供:
  1. TableSpec (列定义 + 标题 + 脚注)
  2. projection function (snapshot → TableData)

本协议固定以下规则 (P0-2 表协议):
  - 表标题: 居中加粗, Heading 3
  - 列顺序: 由 TableSpec.columns 决定, 不可运行时变更
  - 单位: 写在列头括号内, 如 "永久占地 (hm²)"
  - 合计行: has_total_row=True 时必须存在, 加粗
  - 数值格式: int / 1f / 2f / str, 由 TableColumn.fmt 指定
  - 空值: 显示 "—", 不静默输出 0 或空白
  - 数字右对齐, 文字左对齐, 表头居中
"""
from __future__ import annotations

from dataclasses import dataclass, field as dc_field
from enum import Enum
from typing import Any


class TableRenderPolicy(str, Enum):
    """
    P0.5-1: 表格渲染四态 (所有表格必须声明其中之一)

    - RENDER_WITH_VALUES: 数据就绪, 正常渲染表格 + 数据
    - RENDER_WITH_PLACEHOLDER: 数据部分缺失, 渲染表结构 + "—" 占位
    - RENDER_NOT_APPLICABLE: 本项目不适用此表, 渲染标题 + "本项目不涉及此表" 说明
    - SKIP_RENDER: 完全不渲染 (条件不满足, 不保留任何占位)

    规则:
      - 投影函数返回 TableData 时必须设置 render_policy
      - renderer 依据 policy 决定行为, 不做隐式推断
      - RENDER_WITH_PLACEHOLDER 和 RENDER_WITH_VALUES 的区别:
        前者有结构但缺值 (显示 —), 后者完整
    """
    RENDER_WITH_VALUES = "render_with_values"
    RENDER_WITH_PLACEHOLDER = "render_with_placeholder"
    RENDER_NOT_APPLICABLE = "render_not_applicable"
    SKIP_RENDER = "skip_render"

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


@dataclass
class TableColumn:
    key: str          # row dict 中的 key
    header: str       # 显示表头 (不含单位)
    unit: str = ""    # 单位, 如 "hm²", 会显示为 "永久占地 (hm²)"
    align: str = "right"  # "left" | "right" | "center"
    fmt: str = "2f"   # "int" | "1f" | "2f" | "str"


@dataclass
class TableSpec:
    table_id: str           # art.table.* id
    title: str              # 表标题
    columns: list[TableColumn]
    has_total_row: bool = True
    footnote: str = ""
    section_id: str = ""    # 嵌入的 narrative section


@dataclass
class TableData:
    spec: TableSpec
    rows: list[dict]            # [{key: value, ...}, ...]
    total_row: dict | None = None  # 合计行 (key: value)
    render_policy: TableRenderPolicy = TableRenderPolicy.RENDER_WITH_VALUES
    warnings: list[str] = dc_field(default_factory=list)


# ============================================================
# Formatting helpers
# ============================================================

_HEADER_BG = "D9D9D9"


def _format_value(value: Any, fmt: str) -> str:
    """按格式规则格式化单个值 (None 和空字符串都显示 —)"""
    if value is None or value == "":
        return "—"
    try:
        if fmt == "int":
            return str(int(float(value)))
        elif fmt == "1f":
            return f"{float(value):.1f}"
        elif fmt == "2f":
            return f"{float(value):.2f}"
        else:
            return str(value)
    except (ValueError, TypeError):
        return str(value) if value else "—"


def _set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex, qn("w:val"): "clear",
    })
    shading.append(shading_elem)


def _set_cell(cell, text: str, align: str = "left",
              bold: bool = False, size: int = 9):
    p = cell.paragraphs[0]
    if align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


# ============================================================
# Generic DOCX table renderer
# ============================================================

def render_data_table(doc: Document, table_data: TableData):
    """
    把 TableData 渲染成 DOCX 表格。

    所有 CAN_GENERATE 表格共用此函数。
    遵循 P0-2 表协议 (见 module docstring)。

    P0.5-1 渲染策略:
      RENDER_WITH_VALUES / RENDER_WITH_PLACEHOLDER → 正常渲染表格
      RENDER_NOT_APPLICABLE → 只渲染标题 + "本项目不涉及此表"
      SKIP_RENDER → 完全不输出
    """
    policy = table_data.render_policy

    if policy == TableRenderPolicy.SKIP_RENDER:
        return None

    spec = table_data.spec

    if policy == TableRenderPolicy.RENDER_NOT_APPLICABLE:
        doc.add_heading(spec.title, level=3)
        p = doc.add_paragraph()
        run = p.add_run("本项目不涉及此表。")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True
        return None
    cols = spec.columns
    n_cols = len(cols)

    # 标题
    heading = doc.add_heading(spec.title, level=3)

    # 行数: 表头 + 数据行 + (合计行)
    n_data = len(table_data.rows)
    n_rows = 1 + n_data + (1 if spec.has_total_row and table_data.total_row else 0)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表头行
    header_row = table.rows[0]
    for i, col in enumerate(cols):
        header_text = col.header
        if col.unit:
            header_text += f" ({col.unit})"
        _set_cell(header_row.cells[i], header_text, align="center", bold=True, size=9)
        _set_cell_shading(header_row.cells[i], _HEADER_BG)

    # 数据行
    for r_idx, row_data in enumerate(table_data.rows):
        row = table.rows[1 + r_idx]
        for c_idx, col in enumerate(cols):
            raw = row_data.get(col.key)
            text = _format_value(raw, col.fmt)
            _set_cell(row.cells[c_idx], text, align=col.align, size=9)

    # 合计行
    if spec.has_total_row and table_data.total_row:
        total_idx = 1 + n_data
        total_r = table.rows[total_idx]
        for c_idx, col in enumerate(cols):
            raw = table_data.total_row.get(col.key)
            text = _format_value(raw, col.fmt)
            _set_cell(total_r.cells[c_idx], text, align=col.align, bold=True, size=9)

    # 脚注
    if spec.footnote:
        p = doc.add_paragraph()
        run = p.add_run(spec.footnote)
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.italic = True

    return table
